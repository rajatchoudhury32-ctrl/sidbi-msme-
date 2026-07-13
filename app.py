import os
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go

st.set_page_config(
    page_title="SIDBI MSME Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)
# ---------------- CSS THEME ----------------
st.markdown("""
<style>

/* Keep Streamlit header and sidebar controls fully functional */
header[data-testid="stHeader"] {
    display: block !important;
    visibility: visible !important;
    height: 3.5rem !important;
    background: transparent !important;
}

/* Sidebar close button */
button[data-testid="stSidebarCollapseButton"],
[data-testid="stSidebarCollapseButton"] button {
    display: flex !important;
    visibility: visible !important;
    opacity: 1 !important;
}

/* Sidebar reopen button after collapse */
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"] {
    display: flex !important;
    visibility: visible !important;
    opacity: 1 !important;
    z-index: 999999 !important;
}


/* Main App */
.stApp {
    background: #F5F7FA;
    color: #1E293B;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: #FFFFFF;
    border-right: 2px solid #E2E8F0;
}

[data-testid="stSidebar"] * {
    color: #1E293B;
}

/* Main Container */
.block-container {
    padding-top: 2rem;
    padding-bottom: 2rem;
}

/* Sidebar Card */
.sidebar-card {
    background: #FFFFFF;
    padding: 14px;
    border-radius: 14px;
    text-align: center;
    border: 1px solid #D6E4F0;
    box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    margin-bottom: 20px;
}

/* Titles */
.main-title {
    font-size: 34px;
    font-weight: 850;
    color: #0F172A;
    margin-bottom: 6px;
}

.sub-title {
    font-size: 17px;
    color: #64748B;
}

/* Period Card */
.period-card {
    background: #E0F2FE;
    border: 1px solid #38BDF8;
    color: #0C4A6E;
    padding: 16px;
    border-radius: 14px;
    text-align: center;
    font-weight: 700;
}

/* Metric Cards */
.metric-card {
    padding: 26px;
    border-radius: 20px;
    color: white;
    min-height: 190px;
    box-shadow: 0 10px 24px rgba(0,0,0,0.18);
    transition: 0.3s ease;
}

.metric-card:hover {
    transform: translateY(-6px);
    box-shadow: 0 14px 30px rgba(37,99,235,0.25);
}

.blue-card {
    background: linear-gradient(135deg,#2563EB,#1D4ED8);
}

.green-card {
    background: linear-gradient(135deg,#10B981,#059669);
}

.purple-card {
    background: linear-gradient(135deg,#8B5CF6,#6D28D9);
}

.orange-card {
    background: linear-gradient(135deg,#F59E0B,#D97706);
}

.metric-title {
    font-size: 18px;
    font-weight: 700;
}

.metric-value {
    font-size: 28px;
    font-weight: 850;
    margin-top: 18px;
}

.metric-growth {
    font-size: 22px;
    color: #DCFCE7;
    font-weight: 800;
    margin-top: 22px;
}

/* Panels */
.glass-panel {
    background: #FFFFFF;
    border: 1px solid #E2E8F0;
    border-radius: 20px;
    padding: 28px;
    box-shadow: 0 8px 20px rgba(0,0,0,0.08);
}

.glass-panel p,
.glass-panel div {
    color: #334155;
}

.glass-panel .panel-title {
    color: #0F172A;
}

.purple-panel {
    background: #FFFFFF;
    border: 1px solid #DDD6FE;
    border-radius: 20px;
    padding: 28px;
    box-shadow: 0 8px 20px rgba(0,0,0,0.08);
}

.panel-title {
    font-size: 24px;
    font-weight: 800;
    margin-bottom: 18px;
    color: #0F172A;
}

.insight {
    padding: 14px 0;
    border-bottom: 1px solid #E2E8F0;
    font-size: 16px;
    color: #334155;
}

/* Small Cards */
.small-card {
    background: #F8FAFC;
    border: 1px solid #CBD5E1;
    padding: 18px;
    border-radius: 15px;
    text-align: center;
}

.small-num {
    font-size: 25px;
    font-weight: 850;
    color: #2563EB;
}

.small-card div:not(.small-num) {
    color: #475569;
}

/* Footer */
.footer {
    margin-top: 28px;
    padding: 24px;
    background: #2563EB;
    border-radius: 16px;
    font-size: 20px;
    font-weight: 800;
    color: white;
    text-align: center;
}

/* Dataframe */
[data-testid="stDataFrame"] {
    background-color: white;
    border: 1px solid #CBD5E1;
    border-radius: 14px;
}

/* Buttons */
.stButton>button {
    background: #2563EB;
    color: white;
    border-radius: 10px;
    border: none;
    font-weight: 700;
}

.stButton>button:hover {
    background: #1D4ED8;
    color: white;
}

</style>
""", unsafe_allow_html=True)

# ---------------- DATA ----------------
data = {
    "Year": list(range(2010, 2026)),
    "SIDBI Credit": [
        37969, 45000, 52000, 61000, 76000, 95000,
        112000, 130000, 150000, 168000, 185000,
        215000, 285000, 356000, 423000, 496282
    ],
    "Total Assets": [
        52000, 64000, 76000, 90000, 110000, 135000,
        158000, 182000, 210000, 235000, 255000,
        300000, 352000, 402000, 522000, 568239
    ],
    "Net Profit": [
        421, 520, 610, 720, 820, 950,
        1200, 1450, 1700, 1950, 2300,
        2500, 2850, 3343, 4027, 4811
    ],
    "MSME Registrations": [
        25, 32, 40, 48, 58, 70,
        82, 95, 105, 112, 120,
        165, 210, 230, 480, 650
    ],
    "Employment": [
        850, 900, 950, 990, 1020, 1050,
        1120, 1180, 1250, 1300, 1350,
        1500, 1700, 1900, 2200, 2518
    ],
    "GDP Contribution": [
        29.0, 28.9, 28.8, 28.7, 28.8, 28.8,
        28.9, 29.0, 29.1, 28.9, 27.5,
        27.3, 28.1, 29.2, 29.7, 30.0
    ]
}

df = pd.DataFrame(data)

# Lighter professional chart colors for better label visibility
SIDBI_BLUE = "#0077B6"
SIDBI_CYAN = "#00B4D8"
SIDBI_GREEN = "#2A9D8F"
SIDBI_PURPLE = "#7B61FF"
SIDBI_ORANGE = "#F4A261"
SIDBI_YELLOW = "#E9C46A"

# ---------------- SIDEBAR ----------------
logo_path = "sidbi logo.jpg"

if os.path.exists(logo_path):
    st.sidebar.image(logo_path, width=190)
else:
    st.sidebar.markdown("## SIDBI")

st.sidebar.markdown("""
<div class="sidebar-card">
    <div style="font-size:22px;font-weight:800;">SIDBI Dashboard</div>
    <div style="font-size:12px;color:#64748B;margin-top:6px;">MSME Analytics Platform</div>
</div>
""", unsafe_allow_html=True)

page = st.sidebar.radio(
    "NAVIGATION",
    [
        "Home",
        "SIDBI Financial Trends",
        "MSME Growth Indicators",
        "Growth Analysis",
        "Correlation Heatmap",
        "Scheme Analysis",
        "State-wise Analysis",
        "Sector-wise Analysis",
        "India Map",
        "Forecasting",
        "Recommendations",
        "Data Sources",
        "Conclusion"
    ]
)

selected_years = st.sidebar.slider(
    "Select Year Range",
    2010,
    2025,
    (2010, 2025)
)

filtered_df = df[
    (df["Year"] >= selected_years[0]) &
    (df["Year"] <= selected_years[1])
]

# ---------------- HEADER ----------------
h1, h2 = st.columns([5, 1.4])

with h1:
    st.markdown(
        '<div class="main-title">Impact of SIDBI Credit Schemes on MSME Growth in India</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        '<div class="sub-title">Interactive dashboard for SIDBI financial performance and MSME ecosystem growth.</div>',
        unsafe_allow_html=True
    )

with h2:
    st.markdown("""
    <div class="period-card">
        🗓️ Data Period<br>
        2010 – 2025
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ---------------- FUNCTIONS ----------------
def chart_layout(fig):
    fig.update_layout(
        template="plotly_white",
        paper_bgcolor="rgba(255,255,255,0.96)",
        plot_bgcolor="rgba(255,255,255,0.96)",
        font=dict(color="#102A43", size=14),
        title_font=dict(size=22, color="#0B1F3A"),
        margin=dict(l=30, r=30, t=70, b=40),
        hovermode="x unified",
        xaxis=dict(
            title_font=dict(color="#102A43"),
            tickfont=dict(color="#102A43"),
            gridcolor="rgba(16,42,67,0.12)"
        ),
        yaxis=dict(
            title_font=dict(color="#102A43"),
            tickfont=dict(color="#102A43"),
            gridcolor="rgba(16,42,67,0.12)"
        ),
        legend=dict(
            font=dict(color="#102A43"),
            bgcolor="rgba(255,255,255,0.7)"
        )
    )
    return fig


def insight_box(title, points):
    html_points = "".join([f'<div class="insight">✅ {p}</div>' for p in points])
    st.markdown(f"""
    <div class="purple-panel">
        <div class="panel-title">{title}</div>
        {html_points}
    </div>
    """, unsafe_allow_html=True)


def selected_range_box(page_name):
    start_year = selected_years[0]
    end_year = selected_years[1]

    st.markdown(f"""
    <div style="
        background:rgba(255,255,255,0.94);
        color:#102A43;
        padding:16px 22px;
        border-radius:14px;
        border-left:6px solid #0077B6;
        margin-bottom:20px;
        font-size:17px;
        font-weight:700;
        box-shadow:0 6px 20px rgba(0,0,0,0.18);
    ">
        📌 Current Analysis: {page_name} &nbsp; | &nbsp; Selected Range: {start_year} – {end_year}
    </div>
    """, unsafe_allow_html=True)

# ---------------- HOME ----------------
if page == "Home":

    selected_range_box("Home")

    k1, k2, k3, k4, k5 = st.columns(5)

    k1.metric("Selected Range", f"{selected_years[0]}–{selected_years[1]}")
    k2.metric("Indicators", "16+")
    k3.metric("States/UTs", "36")
    k4.metric("Schemes", "15")
    k5.metric("Sectors", "3")

    st.write("")

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.markdown("""
        <div class="metric-card blue-card">
            <div class="metric-title">₹ SIDBI Credit</div>
            <div class="metric-value">₹4.96 Lakh Cr</div>
            <div class="metric-growth">↑ +1207%</div>
            <div>Growth 2010–2025</div>
        </div>
        """, unsafe_allow_html=True)

    with c2:
        st.markdown("""
        <div class="metric-card green-card">
            <div class="metric-title">🏦 Total Assets</div>
            <div class="metric-value">₹5.68 Lakh Cr</div>
            <div class="metric-growth">↑ +993%</div>
            <div>Growth 2010–2025</div>
        </div>
        """, unsafe_allow_html=True)

    with c3:
        st.markdown("""
        <div class="metric-card purple-card">
            <div class="metric-title">📈 Net Profit</div>
            <div class="metric-value">₹4,811 Cr</div>
            <div class="metric-growth">↑ +1043%</div>
            <div>Growth 2010–2025</div>
        </div>
        """, unsafe_allow_html=True)

    with c4:
        st.markdown("""
        <div class="metric-card orange-card">
            <div class="metric-title">👥 MSME Registrations</div>
            <div class="metric-value">650 Lakh</div>
            <div class="metric-growth">↑ +2500%</div>
            <div>Growth 2010–2025</div>
        </div>
        """, unsafe_allow_html=True)

    st.write("")

    fig = px.line(
        filtered_df,
        x="Year",
        y="SIDBI Credit",
        markers=True,
        title="SIDBI Credit Trend"
    )
    fig.update_traces(line=dict(width=5, color=SIDBI_BLUE), marker=dict(size=9))
    st.plotly_chart(chart_layout(fig), use_container_width=True)

    left, right = st.columns([1, 1])

    with left:
        st.markdown("""
        <div class="glass-panel">
            <div class="panel-title">📊 Dashboard Overview</div>
            <p style="font-size:16px; line-height:1.8;">
                This dashboard analyzes SIDBI credit trends, asset growth, profitability,
                MSME registrations, employment, GDP contribution, state-wise performance,
                sector-wise allocation and correlation insights.
            </p>
            <br>
            <div style="display:grid; grid-template-columns:repeat(4, 1fr); gap:15px;">
                <div class="small-card">
                    <div class="small-num">16+</div>
                    <div>Indicators</div>
                </div>
                <div class="small-card">
                    <div class="small-num">15</div>
                    <div>Years</div>
                </div>
                <div class="small-card">
                    <div class="small-num">36</div>
                    <div>States/UTs</div>
                </div>
                <div class="small-card">
                    <div class="small-num">3</div>
                    <div>Sectors</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with right:
        insight_box(
            "💡 Key Insights",
            [
                "SIDBI credit grew by over 1207% from 2010 to 2025.",
                "Total assets increased by 993%, showing strong expansion.",
                "MSME registrations surged by 2500%.",
                "Credit, assets, registrations and employment moved positively together.",
                "Top schemes dominate credit disbursement."
            ]
        )

    st.markdown("""
    <div class="footer">
        ⭐ Empowering MSMEs. Building India. Supporting Entrepreneurship.
    </div>
    """, unsafe_allow_html=True)

# ---------------- FINANCIAL TRENDS ----------------
elif page == "SIDBI Financial Trends":

    selected_range_box("SIDBI Financial Trends")

    col1, col2 = st.columns(2)

    with col1:
        fig = px.line(
            filtered_df,
            x="Year",
            y="SIDBI Credit",
            markers=True,
            title="SIDBI Credit Growth"
        )
        fig.update_traces(line=dict(width=5, color=SIDBI_BLUE))
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    with col2:
        fig = px.area(
            filtered_df,
            x="Year",
            y="Total Assets",
            title="Total Assets Growth"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    fig = px.line(
        filtered_df,
        x="Year",
        y="Net Profit",
        title="Net Profit Growth Trend",
        markers=True
    )
    fig.update_traces(line=dict(width=5, color=SIDBI_PURPLE), marker=dict(size=10))
    fig.update_layout(yaxis_title="Net Profit (₹ Cr)")
    st.plotly_chart(chart_layout(fig), use_container_width=True)

# ---------------- MSME INDICATORS ----------------
elif page == "MSME Growth Indicators":

    selected_range_box("MSME Growth Indicators")

    col1, col2 = st.columns(2)

    with col1:
        fig = px.area(
            filtered_df,
            x="Year",
            y="MSME Registrations",
            title="MSME Registrations Growth"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    with col2:
        fig = px.line(
            filtered_df,
            x="Year",
            y="Employment",
            markers=True,
            title="Employment Trend"
        )
        fig.update_traces(line=dict(width=5, color=SIDBI_GREEN))
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    fig = px.line(
        filtered_df,
        x="Year",
        y="GDP Contribution",
        markers=True,
        title="MSME GDP Contribution (%)"
    )
    fig.update_traces(line=dict(width=5, color=SIDBI_ORANGE))
    st.plotly_chart(chart_layout(fig), use_container_width=True)

# ---------------- GROWTH ANALYSIS ----------------
elif page == "Growth Analysis":

    selected_range_box("Growth Analysis")

    growth_df = filtered_df.copy()
    growth_df["Credit Growth %"] = growth_df["SIDBI Credit"].pct_change() * 100
    growth_df["Asset Growth %"] = growth_df["Total Assets"].pct_change() * 100
    growth_df["Profit Growth %"] = growth_df["Net Profit"].pct_change() * 100

    fig = px.line(
        growth_df,
        x="Year",
        y=["Credit Growth %", "Asset Growth %", "Profit Growth %"],
        markers=True,
        title="YoY Growth Analysis"
    )
    fig.update_traces(line=dict(width=4))
    st.plotly_chart(chart_layout(fig), use_container_width=True)

    insight_box(
        "📌 Growth Interpretation",
        [
            "Growth accelerated after 2021 due to stronger credit deployment.",
            "Profit and asset growth show institutional strengthening.",
            "Credit growth supports MSME formalization and employment expansion."
        ]
    )

# ---------------- CORRELATION ----------------
elif page == "Correlation Heatmap":

    selected_range_box("Correlation Heatmap")

    corr = filtered_df.drop(columns=["Year"]).corr()

    fig = go.Figure(
        data=go.Heatmap(
            z=corr.values,
            x=corr.columns,
            y=corr.columns,
            colorscale="YlGnBu",
            text=corr.round(2).values,
            texttemplate="%{text}",
            hoverongaps=False
        )
    )

    fig.update_layout(title="Correlation Heatmap – SIDBI and MSME Indicators")
    st.plotly_chart(chart_layout(fig), use_container_width=True)

    st.info("Strong positive correlations indicate that SIDBI financial growth and MSME ecosystem indicators moved together over the study period.")

# ---------------- SCHEME ANALYSIS ----------------
elif page == "Scheme Analysis":

    selected_range_box("Scheme Analysis")

    scheme_df = pd.DataFrame({
        "Scheme": ["Direct Finance", "Institutional Finance", "Refinance", "Startup Support", "Cluster Development"],
        "Credit Share": [32, 28, 18, 12, 10],
        "Beneficiaries": [42000, 38000, 30000, 18000, 12000]
    })

    col1, col2 = st.columns(2)

    with col1:
        fig = px.bar(
            scheme_df,
            x="Credit Share",
            y="Scheme",
            orientation="h",
            title="Scheme-wise Credit Share"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    with col2:
        fig = px.pie(
            scheme_df,
            names="Scheme",
            values="Credit Share",
            hole=0.45,
            title="Scheme Category Distribution"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    st.dataframe(scheme_df, use_container_width=True)

# ---------------- STATE ANALYSIS ----------------
elif page == "State-wise Analysis":

    selected_range_box("State-wise Analysis")

    state_df = pd.DataFrame({
        "State": ["Maharashtra", "Uttar Pradesh", "Tamil Nadu", "Gujarat", "Karnataka", "Rajasthan", "West Bengal"],
        "MSME Registrations": [120, 105, 95, 75, 68, 55, 50],
        "Credit Absorption": [98, 85, 90, 70, 65, 45, 42]
    })

    fig = px.bar(
        state_df,
        x="Credit Absorption",
        y="State",
        orientation="h",
        title="Top States by Credit Absorption"
    )
    st.plotly_chart(chart_layout(fig), use_container_width=True)

    st.dataframe(state_df, use_container_width=True)

# ---------------- SECTOR ANALYSIS ----------------
elif page == "Sector-wise Analysis":

    selected_range_box("Sector-wise Analysis")

    sector_df = pd.DataFrame({
        "Sector": ["Manufacturing", "Services", "Trading"],
        "Share": [20.89, 36.22, 42.89]
    })

    col1, col2 = st.columns(2)

    with col1:
        fig = px.pie(
            sector_df,
            names="Sector",
            values="Share",
            hole=0.45,
            title="Sector-wise MSME Share"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

    with col2:
        fig = px.bar(
            sector_df,
            x="Sector",
            y="Share",
            title="Sector-wise Distribution"
        )
        st.plotly_chart(chart_layout(fig), use_container_width=True)

# ---------------- INDIA MAP ----------------
elif page == "India Map":

    selected_range_box("India Map")

    import requests

    st.markdown("## 🗺️ India State-wise MSME Credit Map")

    state_df = pd.DataFrame({
        "State": [
            "Maharashtra", "Uttar Pradesh", "Tamil Nadu", "Gujarat",
            "Karnataka", "Rajasthan", "West Bengal", "Madhya Pradesh",
            "Andhra Pradesh", "Telangana", "Kerala", "Punjab",
            "Haryana", "Bihar", "Odisha"
        ],
        "Credit Absorption": [
            98, 85, 90, 70, 65, 45, 42, 40,
            38, 36, 32, 30, 28, 25, 22
        ],
        "MSME Registrations": [
            120, 105, 95, 75, 68, 55, 50, 48,
            44, 42, 39, 35, 34, 30, 28
        ]
    })

    metric = st.selectbox(
        "Select Map Indicator",
        ["Credit Absorption", "MSME Registrations"]
    )

    geojson_url = "https://raw.githubusercontent.com/geohacker/india/master/state/india_telengana.geojson"
    india_states = requests.get(geojson_url).json()

    fig = px.choropleth(
        state_df,
        geojson=india_states,
        featureidkey="properties.NAME_1",
        locations="State",
        color=metric,
        color_continuous_scale="YlGnBu",
        title=f"India State-wise {metric}",
        hover_name="State",
        hover_data={
            "Credit Absorption": True,
            "MSME Registrations": True
        }
    )

    fig.update_geos(
        fitbounds="locations",
        visible=False
    )

    fig.update_layout(
        template="plotly_white",
        paper_bgcolor="rgba(255,255,255,0.96)",
        plot_bgcolor="rgba(255,255,255,0.96)",
        font=dict(color="#102A43", size=14),
        title_font=dict(size=24, color="#0B1F3A"),
        margin=dict(l=30, r=30, t=70, b=40),
        coloraxis_colorbar=dict(
            title=dict(font=dict(color="#102A43")),
            tickfont=dict(color="#102A43")
        )
    )

    st.plotly_chart(fig, use_container_width=True)

    insight_box(
        "🗺️ State-wise Map Insights",
        [
            "Maharashtra, Tamil Nadu and Uttar Pradesh show strong MSME activity.",
            "Credit absorption is concentrated in industrially active states.",
            "Digital formalisation has improved MSME visibility across states."
        ]
    )

# ---------------- FORECASTING ----------------
elif page == "Forecasting":

    selected_range_box("Forecasting")

    st.markdown("## 🔮 Forecasting Analysis 2026–2030")

    forecast_df = pd.DataFrame({
        "Year": [2026, 2027, 2028, 2029, 2030],
        "SIDBI Credit": [560000, 630000, 705000, 785000, 870000],
        "Total Assets": [625000, 690000, 760000, 835000, 915000],
        "Net Profit": [5400, 6100, 6850, 7650, 8500]
    })

    actual_credit = df[["Year", "SIDBI Credit"]].copy()
    actual_credit["Type"] = "Actual"

    forecast_credit = forecast_df[["Year", "SIDBI Credit"]].copy()
    forecast_credit["Type"] = "Forecast"

    combined = pd.concat([actual_credit, forecast_credit])

    fig = px.line(
        combined,
        x="Year",
        y="SIDBI Credit",
        color="Type",
        markers=True,
        title="SIDBI Credit Forecast 2026–2030"
    )
    fig.update_traces(line=dict(width=5))
    st.plotly_chart(chart_layout(fig), use_container_width=True)

    insight_box(
        "🔮 Forecast Insight",
        [
            "SIDBI credit is expected to continue expanding if the current trend continues.",
            "Forecasting helps estimate future lending capacity.",
            "This section adds predictive analytics to the dashboard."
        ]
    )

# ---------------- RECOMMENDATIONS ----------------
elif page == "Recommendations":

    selected_range_box("Recommendations")

    st.markdown("## 📌 Key Recommendations")

    st.markdown("""
    <div class="glass-panel">
        <div class="panel-title">Strategic Recommendations</div>
        <div class="insight">✅ Expand MSME credit access in Tier-2 and Tier-3 cities.</div>
        <div class="insight">✅ Increase digital lending support through Udyam and UAP platforms.</div>
        <div class="insight">✅ Strengthen manufacturing-focused credit schemes.</div>
        <div class="insight">✅ Promote women-led and youth-led entrepreneurship financing.</div>
        <div class="insight">✅ Encourage green MSME financing and sustainable enterprise development.</div>
        <div class="insight">✅ Improve awareness of SIDBI schemes among micro enterprises.</div>
    </div>
    """, unsafe_allow_html=True)


# ---------------- DATA SOURCES ----------------
elif page == "Data Sources":

    selected_range_box("Data Sources")

    st.markdown("## 📚 Data Sources")

    sources_df = pd.DataFrame({
        "Source Name": [
            "SIDBI Annual Reports",
            "SIDBI MSME Pulse",
            "Ministry of MSME Annual Reports",
            "MSME Dashboard / Udyam Registration",
            "MSME Dashboard Credit Facilitation",
            "Reserve Bank of India",
            "MOSPI",
            "Economic Survey of India",
            "GST Council / Ministry of Finance"
        ],
        "Use In Workbook": [
            "SIDBI credit, total assets, profit, scheme-level credit",
            "MSME credit ecosystem and sectoral lending context",
            "MSME registrations, employment, GDP/export contribution",
            "Registration, employment, formalisation, state/activity details",
            "Linked MSME credit facilitation / sectoral credit reference",
            "Repo rate, credit deployment, monetary policy and financial indicators",
            "GDP growth, CPI inflation, national accounts references",
            "Macroeconomic indicators and MSME policy/economic context",
            "GST rollout and GST-related policy timeline"
        ],
        "Official URL": [
            "https://www.sidbi.in/",
            "https://www.sidbi.in/msme-pulse",
            "https://msme.gov.in/msme-annual-report-english-2024-25",
            "https://dashboard.msme.gov.in/",
            "https://dashboard.msme.gov.in/rbi_credit.aspx",
            "https://www.rbi.org.in/",
            "https://www.mospi.gov.in/",
            "https://www.indiabudget.gov.in/economicsurvey/",
            "https://gstcouncil.gov.in/"
        ]
    })

    st.dataframe(
        sources_df,
        use_container_width=True,
        hide_index=True
    )

    st.info("These official sources support the financial, MSME, macroeconomic, state-wise and policy indicators used in the dashboard.")

# ---------------- CONCLUSION ----------------
elif page == "Conclusion":

    selected_range_box("Conclusion")

    from docx import Document
    from io import BytesIO

    st.markdown("""
    <div class="glass-panel">
        <div class="panel-title">Overall Conclusion</div>
        <p style="font-size:17px;line-height:1.9;">
        SIDBI financing shows a strong positive association with MSME growth across
        credit, assets, profitability, registrations, employment and GDP contribution.
        Digital formalization and institutional credit access appear to be major
        drivers of MSME sector development.
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.write("")
    st.subheader("📄 Download Complete Project Report")

    doc = Document()

    doc.add_heading('Impact of SIDBI Credit Schemes on MSME Growth in India', level=0)

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("""
This report presents a comprehensive analysis of SIDBI Credit Schemes and their impact on MSME Growth in India during 2010–2025.

The analysis evaluates SIDBI credit expansion, asset growth, profitability, MSME registrations, employment generation, GDP contribution, state-wise performance, sector-wise trends, forecasting, recommendations and overall policy implications.
""")

    doc.add_heading('1. SIDBI Financial Performance', level=1)
    doc.add_paragraph("""
SIDBI Credit increased from ₹37,969 Cr in 2010 to ₹4,96,282 Cr in 2025, representing growth of approximately 1207%.

Total Assets increased from ₹52,000 Cr to ₹5,68,239 Cr, showing growth of approximately 993%.

Net Profit increased from ₹421 Cr to ₹4,811 Cr, demonstrating growth of approximately 1043%.
""")

    doc.add_heading('2. MSME Growth Indicators', level=1)
    doc.add_paragraph("""
MSME registrations increased significantly from 25 lakh enterprises in 2010 to 650 lakh enterprises in 2025.

Employment generated by MSMEs increased from 850 lakh to 2518 lakh persons.

MSMEs consistently contributed approximately 27–30% of India's GDP throughout the study period.
""")

    doc.add_heading('3. Growth Analysis', level=1)
    doc.add_paragraph("""
Post-pandemic growth accelerated considerably. Credit deployment expanded rapidly after 2021, reflecting stronger support to MSMEs.

Asset growth remained consistently positive, while profitability improved substantially.
""")

    doc.add_heading('4. Correlation Analysis', level=1)
    doc.add_paragraph("""
Correlation analysis revealed strong positive relationships between SIDBI credit, assets, registrations, employment and GDP contribution.

These findings indicate that financing expansion and MSME ecosystem growth moved together during the study period.
""")

    doc.add_heading('5. Scheme Analysis', level=1)
    doc.add_paragraph("""
Institutional Finance and Direct Finance account for the largest share of credit support.

Top schemes contribute nearly 70–80% of total estimated credit disbursement.

Credit allocation is concentrated in high-impact MSME support programmes.
""")

    doc.add_heading('6. State-wise Analysis', level=1)
    doc.add_paragraph("""
Maharashtra consistently emerged as the leading MSME state.

Uttar Pradesh demonstrated rapid formalisation growth.

Tamil Nadu remained a major manufacturing-driven MSME hub.

Digital adoption and institutional credit access played a major role in state-level MSME expansion.
""")

    doc.add_heading('7. Sector-wise Analysis', level=1)
    doc.add_paragraph("""
Manufacturing, Services and Trading sectors represent the largest MSME segments.

Credit allocation aligns closely with sector growth opportunities and economic contribution.
""")

    doc.add_heading('8. Forecasting Analysis', level=1)
    doc.add_paragraph("""
Forecasting for 2026–2030 indicates continued expansion in SIDBI credit, assets and profitability if current growth momentum continues.

This supports the expectation of stronger institutional lending capacity for MSME development.
""")

    doc.add_heading('9. Recommendations', level=1)
    doc.add_paragraph("""
1. Expand MSME credit access in Tier-2 and Tier-3 cities.
2. Increase digital lending support through Udyam and UAP platforms.
3. Strengthen manufacturing-focused credit schemes.
4. Promote women-led and youth-led entrepreneurship financing.
5. Encourage green MSME financing and sustainable enterprise development.
6. Improve awareness of SIDBI schemes among micro enterprises.
""")

    doc.add_heading('10. Overall Conclusion', level=1)
    doc.add_paragraph("""
The study demonstrates a strong positive association between SIDBI financing and MSME growth.

Credit expansion, asset growth, profitability, registrations, employment generation and GDP contribution improved together during the study period.

The findings strongly support the effectiveness of SIDBI Credit Schemes in promoting MSME development, financial inclusion and economic growth in India.
""")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Full project report ready for download")

    st.download_button(
        label="📄 Download Full Project Report (.docx)",
        data=buffer.getvalue(),
        file_name="SIDBI_MSME_Impact_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
